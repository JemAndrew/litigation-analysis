#!/usr/bin/env python3
"""
Document Loader - Extract text from PDFs, Word docs, emails
British English throughout
"""

from pathlib import Path
from typing import Optional
import PyPDF2
from docx import Document
import email
from email import policy


class DocumentLoader:
    """Extract text from various document formats"""
    
    def __init__(self):
        """Initialise document loader"""
        self.supported_extensions = {'.pdf', '.docx', '.doc', '.txt', '.msg', '.eml'}
    
    def extract_text(self, file_path: Path) -> str:
        """
        Extract text from document
        
        Args:
            file_path: Path to document
            
        Returns:
            Extracted text content
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        # Route to appropriate extractor
        if extension == '.pdf':
            return self._extract_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return self._extract_docx(file_path)
        elif extension == '.txt':
            return self._extract_text(file_path)
        elif extension in ['.msg', '.eml']:
            return self._extract_email(file_path)
        else:
            return ""
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF"""
        try:
            text = ""
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                # Limit to first 100 pages for huge PDFs
                max_pages = min(len(reader.pages), 100)
                
                for page_num in range(max_pages):
                    try:
                        page = reader.pages[page_num]
                        text += page.extract_text() + "\n\n"
                    except:
                        continue
            
            return text.strip()
            
        except Exception as e:
            print(f"⚠️  PDF extraction error ({file_path.name}): {e}")
            return ""
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from Word document"""
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            text = "\n\n".join([para.text for para in doc.paragraphs])
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text += "\n" + row_text
            
            return text.strip()
            
        except Exception as e:
            print(f"⚠️  DOCX extraction error ({file_path.name}): {e}")
            return ""
    
    def _extract_text(self, file_path: Path) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            print(f"⚠️  Text extraction error ({file_path.name}): {e}")
            return ""
    
    def _extract_email(self, file_path: Path) -> str:
        """Extract text from email (.msg or .eml)"""
        try:
            with open(file_path, 'rb') as f:
                msg = email.message_from_binary_file(f, policy=policy.default)
            
            # Build email text
            text = f"From: {msg.get('From', '')}\n"
            text += f"To: {msg.get('To', '')}\n"
            text += f"Date: {msg.get('Date', '')}\n"
            text += f"Subject: {msg.get('Subject', '')}\n\n"
            
            # Get body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == 'text/plain':
                        body = part.get_payload(decode=True)
                        if body:
                            text += body.decode('utf-8', errors='ignore')
            else:
                body = msg.get_payload(decode=True)
                if body:
                    text += body.decode('utf-8', errors='ignore')
            
            return text.strip()
            
        except Exception as e:
            print(f"⚠️  Email extraction error ({file_path.name}): {e}")
            return ""
    
    def is_supported(self, file_path: Path) -> bool:
        """Check if file type is supported"""
        return Path(file_path).suffix.lower() in self.supported_extensions